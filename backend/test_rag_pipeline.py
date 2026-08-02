import os
import shutil
import asyncio
from pathlib import Path
from docx import Document as DocxDocument
import fitz  # PyMuPDF

from app.rag.document_loader import DocumentLoader
from app.rag.text_splitter import TextSplitter
from app.rag.embedding_service import EmbeddingService
from app.rag.vector_store import VectorStore
from app.rag.rag_service import RAGService
from app.services.ai_service import AIService

def create_sample_files():
    docs_dir = Path("./documents")
    docs_dir.mkdir(exist_ok=True)

    # 1. Create Sample TXT file
    txt_path = docs_dir / "sample_attendance.txt"
    txt_path.write_text(
        "College Attendance Rules:\n"
        "Students must maintain a minimum of 75% attendance in all courses to be eligible for end semester examinations.\n"
        "If attendance is between 65% and 74%, condonation may be granted on medical grounds with proper documentation.\n"
        "Students with attendance below 65% will be detained and must repeat the course semester.",
        encoding="utf-8"
    )

    # 2. Create Sample DOCX file
    docx_path = docs_dir / "sample_policy.docx"
    doc = DocxDocument()
    doc.add_heading('Campus Library Policy', 0)
    doc.add_paragraph('Library books can be borrowed for a maximum of 14 days.')
    doc.add_paragraph('A fine of Rs 5 per day is charged for late returns.')
    doc.save(docx_path)

    # 3. Create Sample PDF file
    pdf_path = docs_dir / "sample_grading.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Grading System Guidelines:\nMarks >= 90: Grade O (Outstanding)\nMarks 80-89: Grade A+ (Excellent)\nMarks 70-79: Grade A (Very Good)\nMarks 60-69: Grade B+ (Good)")
    doc.save(str(pdf_path))
    doc.close()

    print("Sample test files created successfully.")
    return txt_path, docx_path, pdf_path

async def run_tests():
    print("\n--- Starting Comprehensive RAG Pipeline Verification ---")

    # Step 1: File creation
    txt_path, docx_path, pdf_path = create_sample_files()

    # Step 2: Document Extraction
    loader = DocumentLoader()
    txt_doc = loader.load_single_document(txt_path)
    docx_doc = loader.load_single_document(docx_path)
    pdf_doc = loader.load_single_document(pdf_path)

    print(f"\n[1] Document Extraction:")
    print(f"  - TXT Extracted: {len(txt_doc['content'])} chars")
    print(f"  - DOCX Extracted: {len(docx_doc['content'])} chars")
    print(f"  - PDF Extracted: {len(pdf_doc['content'])} chars")
    assert "75% attendance" in txt_doc['content']
    assert "14 days" in docx_doc['content']
    assert "Grade O" in pdf_doc['content']

    # Step 3: Text Chunking
    splitter = TextSplitter(chunk_size=1000, chunk_overlap=200)
    txt_chunks = splitter.split_single_document(txt_doc)
    print(f"\n[2] Text Chunking (size=1000, overlap=200):")
    print(f"  - Generated {len(txt_chunks)} chunk(s) from TXT")

    # Step 4: Embeddings Generation
    emb_service = EmbeddingService()
    vectors = emb_service.create_embeddings([txt_chunks[0]['content']])
    print(f"\n[3] Embeddings (all-MiniLM-L6-v2):")
    print(f"  - Vector dimension: {len(vectors[0])} (expected 384)")
    assert len(vectors[0]) == 384

    # Step 5: ChromaDB Storage inside vector_db/
    rag_service = RAGService()
    rag_service.process_and_index_file(txt_path)
    rag_service.process_and_index_file(docx_path)
    rag_service.process_and_index_file(pdf_path)

    vstore = VectorStore()
    total_count = vstore.count()
    print(f"\n[4] ChromaDB Vector Store in vector_db/:")
    print(f"  - Total chunks stored in vector_db/: {total_count}")
    assert total_count > 0

    # Step 6: Similarity Search (Top 5 Chunks)
    top_chunks = rag_service.retrieve_context("What is attendance rule?", top_k=5)
    print(f"\n[5] Similarity Search (Top 5 Chunks for 'What is attendance rule?'):")
    print(f"  - Retrieved {len(top_chunks)} chunk(s)")
    for i, c in enumerate(top_chunks, 1):
        print(f"    Chunk {i} [{c['filename']}]: {c['content'][:80]}... (score: {c['score']:.4f})")

    # Step 7: Chat API Flow with Qwen2.5 / Ollama / Fallback LLM
    ai_svc = AIService()
    print("\n[6] Chat API Prompt & Answer Generation:")
    answer = await ai_svc.get_chat_answer("What is attendance rule?")
    print(f"  - Query: 'What is attendance rule?'")
    try:
        print(f"  - Answer:\n{answer}")
    except Exception:
        print(f"  - Answer:\n{answer.encode('ascii', 'ignore').decode('ascii')}")

    # Fallback test
    missing_answer = await ai_svc.get_chat_answer("What is the quantum teleportation protocol of year 3000?")
    print(f"\n[7] Fallback Response Test for missing info:")
    print(f"  - Query: 'What is the quantum teleportation protocol of year 3000?'")
    try:
        print(f"  - Answer:\n{missing_answer}")
    except Exception:
        print(f"  - Answer:\n{missing_answer.encode('ascii', 'ignore').decode('ascii')}")
    assert "I couldn't find this information in the college knowledge base." in missing_answer

    print("\n" + "=" * 50)
    print("[OK] All RAG pipeline tests PASSED successfully!")
    print("=" * 50)

if __name__ == "__main__":
    import asyncio
    asyncio.run(run_tests())
